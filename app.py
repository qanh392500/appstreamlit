import re
import time
import os
import csv
import io
import zipfile
import threading
import queue
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, date
import requests
from bs4 import BeautifulSoup
import streamlit as st

# ── export libs ──────────────────────────────────────────────────────────────
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

st.set_page_config(page_title="VnTax Crawler", page_icon="🏢", layout="wide")

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) "
                  "Chrome/124.0.0.0 Safari/537.36",
    "Accept-Language": "vi,en;q=0.9",
}

PROVINCES = [
    ("An Giang", "an-giang", "133"),
    ("Bắc Ninh", "bac-ninh", "100"),
    ("Cà Mau", "ca-mau", "143"),
    ("Cần Thơ", "can-tho", "138"),
    ("Cao Bằng", "cao-bang", "90"),
    ("Đà Nẵng", "da-nang", "112"),
    ("Đắk Lắk", "dak-lak", "120"),
    ("Điện Biên", "dien-bien", "102"),
    ("Đồng Nai", "dong-nai", "128"),
    ("Đồng Tháp", "dong-thap", "132"),
    ("Gia Lai", "gia-lai", "119"),
    ("Hà Nội", "ha-noi", "81"),
    ("Hà Tĩnh", "ha-tinh", "108"),
    ("Hải Phòng", "hai-phong", "82"),
    ("Hồ Chí Minh", "ho-chi-minh", "122"),
    ("Huế", "hue", "144"),
    ("Hưng Yên", "hung-yen", "84"),
    ("Khánh Hòa", "khanh-hoa", "117"),
    ("Lai Châu", "lai-chau", "103"),
    ("Lâm Đồng", "lam-dong", "123"),
    ("Lạng Sơn", "lang-son", "93"),
    ("Lào Cai", "lao-cai", "91"),
    ("Nghệ An", "nghe-an", "107"),
    ("Ninh Bình", "ninh-binh", "88"),
    ("Phú Thọ", "phu-tho", "97"),
    ("Quảng Ngãi", "quang-ngai", "114"),
    ("Quảng Ninh", "quang-ninh", "101"),
    ("Quảng Trị", "quang-tri", "110"),
    ("Sơn La", "son-la", "104"),
    ("Tây Ninh", "tay-ninh", "126"),
    ("Thái Nguyên", "thai-nguyen", "96"),
    ("Thanh Hóa", "thanh-hoa", "106"),
    ("Tuyên Quang", "tuyen-quang", "94"),
    ("Vĩnh Long", "vinh-long", "135"),
]

DETAIL_RE = re.compile(r"^/\d+(-\d+)?-[a-z0-9-]+$")
OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
os.makedirs(OUT_DIR, exist_ok=True)

FIELDS = ["Mã số thuế", "Tên công ty", "Ngày hoạt động", "Tình trạng", "Địa chỉ"]

# ── helpers ───────────────────────────────────────────────────────────────────

def make_session():
    s = requests.Session()
    s.headers.update(HEADERS)
    return s

def fetch(sess, url, stop_event, retries=3):
    for i in range(retries):
        if stop_event.is_set():
            return None
        try:
            r = sess.get(url, timeout=30)
            if r.status_code == 200:
                return r.text
        except Exception:
            pass
        time.sleep(2 * (i + 1))
    return None

def parse_list_page(html):
    s = BeautifulSoup(html, "html.parser")
    items = []
    for a in s.find_all("a", href=DETAIL_RE):
        h3 = a.find("h3")
        if not h3:
            continue
        text = a.get_text("\n", strip=True)
        m = re.search(r"Mã số thuế:\s*([\d\-]+)", text)
        mst = m.group(1) if m else ""
        m = re.search(r"Địa chỉ:\s*(.+)", text)
        addr = m.group(1).strip() if m else ""
        items.append({
            "name": h3.get_text(strip=True),
            "mst": mst,
            "addr": addr,
            "detail_url": "https://vntax.net" + a["href"],
        })
    cur = re.search(r'class="[^"]*bg-blue-600[^"]*"[^>]*>\s*(\d+)\s*<', html)
    return items, (int(cur.group(1)) if cur else None)

def parse_detail(html):
    if not html:
        return "", ""
    s = BeautifulSoup(html, "html.parser")
    lines = [l.strip() for l in s.get_text("\n", strip=True).split("\n") if l.strip()]
    ngay_hd = tinh_trang = ""
    for i, line in enumerate(lines):
        if line == "Ngày hoạt động" and i + 1 < len(lines):
            val = lines[i + 1]
            if re.match(r"\d{1,2}/\d{1,2}/\d{4}", val):
                ngay_hd = val
        if line == "Tình trạng" and i + 1 < len(lines):
            tinh_trang = lines[i + 1]
    return ngay_hd, tinh_trang

def parse_date(s):
    m = re.match(r"(\d{1,2})/(\d{1,2})/(\d{4})", s or "")
    if m:
        try:
            return date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
        except Exception:
            pass
    return None

def sort_by_date_desc(rows):
    # mới → cũ, không có ngày xuống cuối
    has = [r for r in rows if parse_date(r["Ngày hoạt động"])]
    no  = [r for r in rows if not parse_date(r["Ngày hoạt động"])]
    has.sort(key=lambda r: parse_date(r["Ngày hoạt động"]), reverse=True)
    return has + no

def in_range(row, date_from, date_to):
    if not date_from and not date_to:
        return True
    d = parse_date(row["Ngày hoạt động"])
    if d is None:
        return True  # không có ngày → vẫn giữ
    if date_from and d < date_from:
        return False
    if date_to and d > date_to:
        return False
    return True

# ── export ────────────────────────────────────────────────────────────────────

def to_excel(rows, title):
    buf = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31]
    header_fill = PatternFill("solid", fgColor="2B58C5")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    for ci, field in enumerate(FIELDS, 1):
        cell = ws.cell(row=1, column=ci, value=field)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    for ri, row in enumerate(rows, 2):
        for ci, field in enumerate(FIELDS, 1):
            ws.cell(row=ri, column=ci, value=row[field])
    col_widths = [18, 50, 18, 20, 60]
    for ci, w in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = w
    wb.save(buf)
    return buf.getvalue()

def to_txt(rows, title):
    buf = io.StringIO()
    buf.write(f"{title}\n{'='*80}\n\n")
    for r in rows:
        buf.write(f"MST        : {r['Mã số thuế']}\n")
        buf.write(f"Tên CT     : {r['Tên công ty']}\n")
        buf.write(f"Ngày HD    : {r['Ngày hoạt động']}\n")
        buf.write(f"Tình trạng : {r['Tình trạng']}\n")
        buf.write(f"Địa chỉ    : {r['Địa chỉ']}\n")
        buf.write("-"*80 + "\n")
    return buf.getvalue().encode("utf-8-sig")

def to_docx(rows, title):
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading(title, 0)
    table = doc.add_table(rows=1, cols=len(FIELDS))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, f in enumerate(FIELDS):
        hdr[i].text = f
        hdr[i].paragraphs[0].runs[0].font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, f in enumerate(FIELDS):
            cells[i].text = row[f]
    doc.save(buf)
    return buf.getvalue()

def to_pdf(rows, title):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=20, rightMargin=20,
                            topMargin=30, bottomMargin=30)

    # Try to register a Vietnamese-capable font
    font_name = "Helvetica"
    for font_path in [
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/times.ttf",
    ]:
        if os.path.exists(font_path):
            try:
                fname = os.path.splitext(os.path.basename(font_path))[0]
                pdfmetrics.registerFont(TTFont(fname, font_path))
                font_name = fname
                break
            except Exception:
                pass

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", fontName=font_name, fontSize=14, spaceAfter=12, alignment=1)
    cell_style = ParagraphStyle("cell", fontName=font_name, fontSize=7, leading=10)

    story = [Paragraph(title, title_style), Spacer(1, 8)]

    col_widths = [80, 180, 70, 80, 150]
    data = [[Paragraph(f"<b>{h}</b>", cell_style) for h in FIELDS]]
    for row in rows:
        data.append([Paragraph(row[f] or "", cell_style) for f in FIELDS])

    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2B58C5")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("GRID",       (0, 0), (-1, -1), 0.4, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fc")]),
        ("VALIGN",     (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    doc.build(story)
    return buf.getvalue()

# ── crawler ───────────────────────────────────────────────────────────────────

def fetch_detail_worker(args):
    item, stop_event = args
    if stop_event.is_set():
        return item, "", ""
    sess = make_session()
    html = fetch(sess, item["detail_url"], stop_event)
    ngay_hd, tinh_trang = parse_detail(html)
    return item, ngay_hd, tinh_trang

def crawl(selected_provinces, log_q, progress_q, stop_event, threads, date_from, date_to, max_pages_default=999, max_pages_per_prov=None):
    total_prov = len(selected_provinces)

    for prov_idx, (name, slug, pid) in enumerate(selected_provinces, 1):
        if stop_event.is_set():
            log_q.put(("log", "⛔ Đã dừng theo yêu cầu."))
            break

        base_url = f"https://vntax.net/tra-cuu-ma-so-thue-theo-tinh/{slug}-{pid}"
        log_q.put(("log", f"{'='*48}"))
        max_p = (max_pages_per_prov or {}).get(name) or max_pages_default
        log_q.put(("log", f"[{prov_idx}/{total_prov}] **{name}** (tối đa {max_p} trang)"))
        all_rows = []

        for page in range(1, max_p + 1):
            if stop_event.is_set():
                log_q.put(("log", "⛔ Dừng giữa chừng."))
                break

            url = base_url if page == 1 else f"{base_url}?page={page}"
            html = fetch(make_session(), url, stop_event)
            if not html:
                log_q.put(("log", f"  ⚠️ Trang {page} lỗi, bỏ qua"))
                break
            items, cur = parse_list_page(html)
            if not items:
                log_q.put(("log", f"  Trang {page}: không có dữ liệu, dừng"))
                break
            if cur is not None and cur != page:
                log_q.put(("log", f"  Trang {page}: hết dữ liệu, dừng"))
                break

            total_items = len(items)
            log_q.put(("log", f"  Trang {page}: {total_items} công ty — đang lấy chi tiết..."))
            log_q.put(("page_progress", 0, total_items, page))

            page_rows = []
            done = 0

            with ThreadPoolExecutor(max_workers=threads) as ex:
                futures = [ex.submit(fetch_detail_worker, (item, stop_event)) for item in items]
                for future in as_completed(futures):
                    if stop_event.is_set():
                        break
                    item, ngay_hd, tinh_trang = future.result()
                    page_rows.append({
                        "Mã số thuế": item["mst"],
                        "Tên công ty": item["name"],
                        "Ngày hoạt động": ngay_hd,
                        "Tình trạng": tinh_trang,
                        "Địa chỉ": item["addr"],
                        "_page": page,
                        "_order": len(all_rows) + len(page_rows),
                    })
                    done += 1
                    log_q.put(("page_progress", done, total_items, page))

            all_rows.extend(page_rows)
            log_q.put(("log", f"  Trang {page}: ✅ xong ({len(page_rows)} công ty, tổng: {len(all_rows)})"))
            time.sleep(0.3)

        filtered_page  = [r for r in all_rows   if in_range(r, date_from, date_to)]
        filtered_date  = sort_by_date_desc(filtered_page[:])  # mới → cũ
        # giữ thứ tự trang (sort by _order)
        filtered_page.sort(key=lambda r: r["_order"])

        log_q.put(("log", f"  Lọc ngày: {len(all_rows)} → {len(filtered_date)} công ty"))

        safe_name = name.replace(" ", "_")
        if filtered_date:
            out_path = os.path.join(OUT_DIR, f"{safe_name}.pdf")
            pdf_bytes = to_pdf(filtered_date, f"Danh sách công ty — {name}")
            with open(out_path, "wb") as f:
                f.write(pdf_bytes)
            log_q.put(("log", f"  💾 Lưu {len(filtered_date)} dòng → {out_path}"))
        else:
            log_q.put(("log", "  ⚠️ Không có dữ liệu trong khoảng ngày đã chọn"))
        progress_q.put((name, filtered_date, filtered_page))

    log_q.put(("done",))

# ── session state ─────────────────────────────────────────────────────────────
for k, v in [("running", False), ("logs", []), ("results", {}),
              ("stop_event", None), ("log_q", None), ("progress_q", None),
              ("done_count", 0), ("total_selected", 0), ("stopped", False),
              ("page_prog_text", ""), ("page_prog_val", 0.0),
              ("prov_pages", {p[0]: 3 for p in PROVINCES})]:
    if k not in st.session_state:
        st.session_state[k] = v

# ── UI ────────────────────────────────────────────────────────────────────────
st.title("🏢 VnTax Crawler")
st.caption("Tra cứu mã số thuế theo tỉnh thành — vntax.net")

tab_crawl, tab_filter, tab_lookup = st.tabs(["🔍 Crawl dữ liệu", "📅 Lọc file", "🔎 Tra cứu MST"])

# ════════════════════════════════════════════════════════════════
# TAB 1: CRAWL
# ════════════════════════════════════════════════════════════════
with tab_crawl:
    col_left, col_right = st.columns([1, 2])

    with col_left:
        st.subheader("⚙️ Cấu hình")

        # Đa luồng
        st.markdown("**Số luồng crawl chi tiết**")
        threads = st.slider("Luồng", min_value=1, max_value=5, value=3,
                            help="Nhiều luồng hơn = nhanh hơn")

        st.divider()

        # Định dạng xuất
        st.markdown("**Định dạng xuất file**")
        export_fmt = st.radio("Định dạng", ["PDF", "Excel (.xlsx)", "Word (.docx)", "Text (.txt)"],
                              index=0, horizontal=False)

        st.divider()

        # Số trang mặc định + nút áp dụng
        st.markdown("**Số trang tối đa (mặc định cho tất cả tỉnh)**")
        col_pg_in, col_pg_btn = st.columns([2, 1])
        with col_pg_in:
            max_pages_global = st.number_input(
                "Trang", min_value=1, max_value=999, value=3,
                help="999 = crawl hết, tự dừng khi site hết dữ liệu",
                disabled=st.session_state.running,
                label_visibility="collapsed",
            )
        with col_pg_btn:
            if st.button("Áp dụng tất cả", disabled=st.session_state.running):
                for pname, _, pid in PROVINCES:
                    st.session_state.prov_pages[pname] = int(max_pages_global)
                    st.session_state[f"pg_{pid}"] = int(max_pages_global)
                st.rerun()

        st.divider()

        # Chọn tỉnh + override trang
        st.markdown("**Chọn tỉnh thành**")

        chon_tat_ca = st.checkbox("✅ Chọn tất cả", key="chk_all")
        if chon_tat_ca:
            for _, _, pid in PROVINCES:
                st.session_state[f"chk_{pid}"] = True
        elif "chk_all_prev" in st.session_state and st.session_state.chk_all_prev:
            for _, _, pid in PROVINCES:
                st.session_state[f"chk_{pid}"] = False
        st.session_state.chk_all_prev = chon_tat_ca

        selected = []
        max_pages_per_prov = {}
        for pname, slug, pid in PROVINCES:
            if f"pg_{pid}" not in st.session_state:
                st.session_state[f"pg_{pid}"] = st.session_state.prov_pages.get(pname, 3)
            col_chk, col_pg = st.columns([3, 2])
            with col_chk:
                checked = st.checkbox(pname, key=f"chk_{pid}")
            with col_pg:
                st.number_input(
                    "trang", min_value=1, max_value=999,
                    key=f"pg_{pid}",
                    label_visibility="collapsed",
                    disabled=st.session_state.running,
                )
            if checked:
                selected.append((pname, slug, pid))
                max_pages_per_prov[pname] = int(st.session_state.get(f"pg_{pid}", 3))

    with col_right:
        st.subheader("📊 Tiến độ")
        st.info(f"📁 File tự động lưu tại: `{OUT_DIR}`")

        # ── Drain queues mỗi lần rerun ──────────────────────────────────────
        if st.session_state.running:
            log_q = st.session_state.log_q
            progress_q = st.session_state.progress_q
            finished = False
            while True:
                try:
                    msg = log_q.get_nowait()
                    if msg[0] == "done":
                        finished = True
                    elif msg[0] == "log":
                        st.session_state.logs.append(msg[1])
                    elif msg[0] == "page_progress":
                        _, done, total, page = msg
                        st.session_state.page_prog_val = done / total if total else 0
                        st.session_state.page_prog_text = f"Trang {page}: {done}/{total} công ty"
                except queue.Empty:
                    break
            while not progress_q.empty():
                pname, rows_date, rows_page = progress_q.get()
                st.session_state.results[pname] = (rows_date, rows_page)
                st.session_state.done_count += 1
            if finished:
                st.session_state.running = False

        # ── Nút bấm ─────────────────────────────────────────────────────────
        btn_col1, btn_col2, btn_col3 = st.columns([1, 1, 1])
        with btn_col1:
            btn_start = st.button(
                "▶ Bắt đầu crawl",
                disabled=st.session_state.running or len(selected) == 0,
                type="primary",
            )
        with btn_col2:
            btn_stop = st.button(
                "⏹ Dừng",
                disabled=not st.session_state.running,
                type="secondary",
            )
        with btn_col3:
            btn_reset = st.button(
                "🔄 Chạy lại",
                disabled=st.session_state.running,
                type="secondary",
            )

        if btn_stop and st.session_state.stop_event:
            st.session_state.stop_event.set()
            st.session_state.stopped = True

        if btn_reset:
            st.session_state.logs = []
            st.session_state.results = {}
            st.session_state.done_count = 0
            st.session_state.stopped = False
            st.session_state.page_prog_val = 0.0
            st.session_state.page_prog_text = ""
            st.rerun()

        if btn_start and selected and not st.session_state.running:
            stop_event = threading.Event()
            log_q = queue.Queue()
            progress_q = queue.Queue()
            st.session_state.running = True
            st.session_state.stopped = False
            st.session_state.logs = []
            st.session_state.results = {}
            st.session_state.done_count = 0
            st.session_state.total_selected = len(selected)
            st.session_state.stop_event = stop_event
            st.session_state.log_q = log_q
            st.session_state.progress_q = progress_q
            st.session_state.page_prog_val = 0.0
            st.session_state.page_prog_text = ""
            threading.Thread(
                target=crawl,
                args=(selected, log_q, progress_q, stop_event, threads, None, None,
                      999, max_pages_per_prov),
                daemon=True,
            ).start()
            st.rerun()

        # ── Tiến độ ─────────────────────────────────────────────────────────
        total_sel = st.session_state.total_selected
        done_cnt  = st.session_state.done_count
        if total_sel > 0:
            st.progress(done_cnt / total_sel, text=f"Tỉnh: {done_cnt}/{total_sel}")
        if st.session_state.page_prog_text:
            st.progress(st.session_state.page_prog_val, text=st.session_state.page_prog_text)
        if st.session_state.logs:
            st.markdown("\n\n".join(st.session_state.logs[-60:]))
        if st.session_state.running:
            time.sleep(0.5)
            st.rerun()
        elif st.session_state.stopped:
            st.warning(f"Đã dừng. Hoàn thành {done_cnt}/{total_sel} tỉnh.")
        elif total_sel > 0 and done_cnt == total_sel:
            st.success(f"✅ Hoàn thành tất cả {total_sel} tỉnh!")

        # ── Download ────────────────────────────────────────────────────────
        if st.session_state.results:
            st.subheader("⬇ Tải file")
            fmt_map = {
                "PDF": ("pdf", "application/pdf", to_pdf),
                "Excel (.xlsx)": ("xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", to_excel),
                "Word (.docx)": ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", to_docx),
                "Text (.txt)": ("txt", "text/plain", to_txt),
            }
            ext, mime, fn_export = fmt_map[export_fmt]

            def clean(rows):
                return [{k: v for k, v in r.items() if not k.startswith("_")} for r in rows]

            def to_short(rows, title, ext=ext):
                short = [{"Mã số thuế": r["Mã số thuế"], "Ngày hoạt động": r["Ngày hoạt động"]} for r in rows]
                if ext == "pdf":
                    buf = io.BytesIO()
                    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=40, rightMargin=40,
                                            topMargin=30, bottomMargin=30)
                    font_name = "Helvetica"
                    for fp in ["C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/calibri.ttf"]:
                        if os.path.exists(fp):
                            try:
                                fn = os.path.splitext(os.path.basename(fp))[0]
                                pdfmetrics.registerFont(TTFont(fn, fp))
                                font_name = fn; break
                            except Exception:
                                pass
                    ts = ParagraphStyle("t", fontName=font_name, fontSize=13, spaceAfter=10, alignment=1)
                    cs = ParagraphStyle("c", fontName=font_name, fontSize=9, leading=12)
                    story = [Paragraph(title, ts), Spacer(1, 8)]
                    data = [[Paragraph("<b>Mã số thuế</b>", cs), Paragraph("<b>Ngày hoạt động</b>", cs)]]
                    for r in short:
                        data.append([Paragraph(r["Mã số thuế"], cs), Paragraph(r["Ngày hoạt động"], cs)])
                    t = Table(data, colWidths=[180, 120], repeatRows=1)
                    t.setStyle(TableStyle([
                        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#2B58C5")),
                        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                        ("GRID",       (0,0), (-1,-1), 0.4, colors.grey),
                        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f5f7fc")]),
                        ("VALIGN",     (0,0), (-1,-1), "TOP"),
                        ("TOPPADDING", (0,0), (-1,-1), 4),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                    ]))
                    story.append(t); doc.build(story)
                    return buf.getvalue()
                elif ext == "xlsx":
                    buf = io.BytesIO()
                    wb = openpyxl.Workbook(); ws = wb.active; ws.title = title[:31]
                    hf = PatternFill("solid", fgColor="2B58C5"); hfont = Font(bold=True, color="FFFFFF")
                    for ci, f in enumerate(["Mã số thuế", "Ngày hoạt động"], 1):
                        c = ws.cell(row=1, column=ci, value=f); c.fill = hf; c.font = hfont
                    for ri, r in enumerate(short, 2):
                        ws.cell(row=ri, column=1, value=r["Mã số thuế"])
                        ws.cell(row=ri, column=2, value=r["Ngày hoạt động"])
                    ws.column_dimensions["A"].width = 22; ws.column_dimensions["B"].width = 18
                    wb.save(buf); return buf.getvalue()
                elif ext == "docx":
                    buf = io.BytesIO(); doc2 = Document(); doc2.add_heading(title, 0)
                    tbl = doc2.add_table(rows=1, cols=2); tbl.style = "Table Grid"
                    hdr = tbl.rows[0].cells
                    hdr[0].text = "Mã số thuế"; hdr[0].paragraphs[0].runs[0].font.bold = True
                    hdr[1].text = "Ngày hoạt động"; hdr[1].paragraphs[0].runs[0].font.bold = True
                    for r in short:
                        cells = tbl.add_row().cells
                        cells[0].text = r["Mã số thuế"]; cells[1].text = r["Ngày hoạt động"]
                    doc2.save(buf); return buf.getvalue()
                else:
                    buf2 = io.StringIO()
                    buf2.write(f"{title}\n{'='*40}\n\n")
                    for r in short:
                        buf2.write(f"{r['Mã số thuế']}\t{r['Ngày hoạt động']}\n")
                    return buf2.getvalue().encode("utf-8-sig")

            def make_zip(pname, rows_date, rows_page, ext, fn_export, to_short, clean):
                sname = pname.replace(" ", "_")
                buf = io.BytesIO()
                with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    zf.writestr(f"{sname}_day.{ext}",
                        fn_export(clean(rows_date), f"Danh sách — {pname} (mới→cũ)"))
                    zf.writestr(f"{sname}_page.{ext}",
                        fn_export(clean(rows_page), f"Danh sách — {pname} (theo trang)"))
                    zf.writestr(f"{sname}_short_day.{ext}",
                        to_short(rows_date, f"Rút gọn — {pname} (mới→cũ)"))
                    zf.writestr(f"{sname}_short_page.{ext}",
                        to_short(rows_page, f"Rút gọn — {pname} (theo trang)"))
                return buf.getvalue()

            for pname, (rows_date, rows_page) in st.session_state.results.items():
                if not rows_date and not rows_page:
                    st.caption(f"⚠️ {pname}: không có dữ liệu")
                    continue
                st.markdown(f"**{pname}** ({len(rows_date)} công ty)")
                c1, c2, c3, c4, c5 = st.columns(5)
                sname = pname.replace(" ", "_")
                with c1:
                    st.download_button("⬇ Đầy đủ · mới→cũ",
                        data=fn_export(clean(rows_date), f"Danh sách — {pname} (mới→cũ)"),
                        file_name=f"{sname}_day.{ext}", mime=mime, key=f"dl1_{pname}")
                with c2:
                    st.download_button("⬇ Đầy đủ · theo trang",
                        data=fn_export(clean(rows_page), f"Danh sách — {pname} (theo trang)"),
                        file_name=f"{sname}_page.{ext}", mime=mime, key=f"dl2_{pname}")
                with c3:
                    st.download_button("⬇ Rút gọn · mới→cũ",
                        data=to_short(rows_date, f"Rút gọn — {pname} (mới→cũ)"),
                        file_name=f"{sname}_short_day.{ext}", mime=mime, key=f"dl3_{pname}")
                with c4:
                    st.download_button("⬇ Rút gọn · theo trang",
                        data=to_short(rows_page, f"Rút gọn — {pname} (theo trang)"),
                        file_name=f"{sname}_short_page.{ext}", mime=mime, key=f"dl4_{pname}")
                with c5:
                    st.download_button("⬇ Tất cả (ZIP)",
                        data=make_zip(pname, rows_date, rows_page, ext, fn_export, to_short, clean),
                        file_name=f"{sname}_all.zip", mime="application/zip", key=f"dl5_{pname}")

# ════════════════════════════════════════════════════════════════
# TAB 2: LỌC FILE
# ════════════════════════════════════════════════════════════════
with tab_filter:
    st.subheader("📅 Lọc theo Ngày hoạt động")
    st.caption("Upload file đã crawl, chọn khoảng ngày, tải file đã lọc.")

    col_f1, col_f2 = st.columns([1, 2])

    with col_f1:
        st.markdown("**Upload file**")
        uploaded = st.file_uploader(
            "Chọn file (Excel hoặc CSV)",
            type=["xlsx", "csv"],
            help="Upload file Excel/CSV đã crawl từ Tab Crawl"
        )

        st.divider()
        st.markdown("**Lọc theo Ngày hoạt động** *(bao gồm cả 2 đầu mút)*")
        c1, c2 = st.columns(2)
        with c1:
            f_date_from = st.date_input("Từ ngày", value=None, format="DD/MM/YYYY", key="f_from")
        with c2:
            f_date_to = st.date_input("Đến ngày", value=None, format="DD/MM/YYYY", key="f_to")

        if f_date_from and f_date_to and f_date_from > f_date_to:
            st.error("Từ ngày phải ≤ Đến ngày")
            f_date_from = f_date_to = None

        st.divider()
        st.markdown("**Định dạng xuất**")
        f_fmt = st.radio("Định dạng", ["PDF", "Excel (.xlsx)", "Word (.docx)", "Text (.txt)"],
                         index=0, key="f_fmt")

        btn_filter = st.button("🔍 Lọc", type="primary", disabled=uploaded is None)

    with col_f2:
        if uploaded is not None:
            # Đọc file
            try:
                if uploaded.name.endswith(".csv"):
                    import pandas as pd
                    df_up = pd.read_csv(uploaded)
                else:
                    import pandas as pd
                    df_up = pd.read_excel(uploaded)

                # Chuẩn hoá cột
                df_up.columns = [c.strip() for c in df_up.columns]
                rows_up = df_up.fillna("").to_dict("records")
                # Đảm bảo các cột cần thiết tồn tại
                for r in rows_up:
                    if "Mã số thuế" not in r: r["Mã số thuế"] = ""
                    if "Ngày hoạt động" not in r: r["Ngày hoạt động"] = ""
                    if "Tên công ty" not in r: r["Tên công ty"] = ""
                    if "Tình trạng" not in r: r["Tình trạng"] = ""
                    if "Địa chỉ" not in r: r["Địa chỉ"] = ""

                st.success(f"Đã đọc {len(rows_up)} dòng từ **{uploaded.name}**")

                # Preview
                st.dataframe(df_up.head(10), use_container_width=True)

                if btn_filter:
                    filtered_up = [r for r in rows_up if in_range(r, f_date_from, f_date_to)]
                    filtered_up = sort_by_date_desc(filtered_up)

                    st.info(f"Kết quả lọc: **{len(filtered_up)}** / {len(rows_up)} dòng")

                    if filtered_up:
                        f_fmt_map = {
                            "PDF": ("pdf", "application/pdf", to_pdf),
                            "Excel (.xlsx)": ("xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", to_excel),
                            "Word (.docx)": ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", to_docx),
                            "Text (.txt)": ("txt", "text/plain", to_txt),
                        }
                        f_ext, f_mime, f_fn = f_fmt_map[f_fmt]
                        title = f"Kết quả lọc — {uploaded.name}"
                        data = f_fn(filtered_up, title)
                        base = os.path.splitext(uploaded.name)[0]
                        st.download_button(
                            label=f"⬇ Tải kết quả ({len(filtered_up)} dòng) — {f_ext.upper()}",
                            data=data,
                            file_name=f"{base}_filtered.{f_ext}",
                            mime=f_mime,
                            key="dl_filtered",
                        )
                    else:
                        st.warning("Không có dòng nào trong khoảng ngày đã chọn.")

            except Exception as e:
                st.error(f"Lỗi đọc file: {e}")
        else:
            st.info("Upload file Excel hoặc CSV để bắt đầu lọc.")

# ════════════════════════════════════════════════════════════════
# TAB 3: TRA CỨU MST
# ════════════════════════════════════════════════════════════════

def lookup_mst(mst_raw):
    mst = mst_raw.strip()
    sess = make_session()
    stop = threading.Event()

    # Dùng search để lấy URL đúng (tránh redirect về trang mẹ)
    search_html = fetch(sess, f"https://vntax.net/search?q={mst}", stop)
    target_url = f"https://vntax.net/{mst}"
    if search_html:
        soup_s = BeautifulSoup(search_html, "html.parser")
        # URL sau redirect của search chính là trang đúng
        # Hoặc tìm link đầu tiên khớp MST
        for a in soup_s.find_all("a", href=True):
            href = a["href"]
            if href.startswith(f"/{mst}-") or href == f"/{mst}":
                target_url = "https://vntax.net" + href
                break

    # Nếu search redirect thẳng về trang chi tiết
    try:
        r = sess.get(f"https://vntax.net/search?q={mst}", timeout=30, allow_redirects=True)
        if f"/{mst}" in r.url and r.url != f"https://vntax.net/search?q={mst}":
            target_url = r.url
    except Exception:
        pass

    html = fetch(sess, target_url, stop)
    if not html:
        return None, "Không thể kết nối hoặc không tìm thấy MST."
    soup = BeautifulSoup(html, "html.parser")

    # Tên công ty từ h1
    h1 = soup.find("h1")
    ten_ct = h1.get_text(strip=True) if h1 else ""

    # Parse các trường từ grid rows (sm:col-span-3 = label, sm:col-span-9 = value)
    info = {}
    rows = soup.select("div.grid.grid-cols-1.sm\:grid-cols-12")
    for row in rows:
        label_div = row.select_one("div.sm\:col-span-3")
        value_div = row.select_one("div.sm\:col-span-9")
        if not label_div or not value_div:
            continue
        # Bỏ icon span trước khi lấy text
        for span in label_div.find_all("span"):
            span.decompose()
        label = label_div.get_text(strip=True)
        # Số điện thoại: lấy data-full thay vì text bị mask
        phone_span = value_div.find("span", attrs={"data-full": True})
        if phone_span:
            value = phone_span["data-full"]
        else:
            value = value_div.get_text(strip=True)
            value = re.sub(r'Hiển thị\s*$', '', value).strip()
        if label:
            info[label] = value

    # Bảng ngành nghề kinh doanh
    nganh_nghe = []
    for div in soup.select("div.grid-cols-12"):
        ma_div = div.select_one(".col-span-2")
        nganh_div = div.select_one(".col-span-10")
        if ma_div and nganh_div:
            ma = ma_div.get_text(strip=True)
            if ma.isdigit() and len(ma) == 4:
                nganh_nghe.append({"Mã": ma, "Ngành": nganh_div.get_text(" ", strip=True)})

    return {"ten": ten_ct, "info": info, "nganh_nghe": nganh_nghe}, None


def lookup_to_pdf(ten, info, nganh_nghe):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=40, rightMargin=40,
                            topMargin=40, bottomMargin=40)

    font_name = "Helvetica"
    for fp in [
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/times.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    ]:
        if os.path.exists(fp):
            try:
                fn = os.path.splitext(os.path.basename(fp))[0]
                pdfmetrics.registerFont(TTFont(fn, fp))
                font_name = fn
                break
            except Exception:
                pass

    title_style  = ParagraphStyle("t", fontName=font_name, fontSize=13, spaceAfter=14,
                                   textColor=colors.HexColor("#1a56db"), leading=18)
    label_style  = ParagraphStyle("l", fontName=font_name, fontSize=9,
                                   textColor=colors.HexColor("#6b7280"))
    value_style  = ParagraphStyle("v", fontName=font_name, fontSize=10, leading=14)
    head_style   = ParagraphStyle("h", fontName=font_name, fontSize=10, spaceAfter=6,
                                   textColor=colors.white)
    cell_style   = ParagraphStyle("c", fontName=font_name, fontSize=9, leading=12)

    story = [Paragraph(ten, title_style)]

    for field, val in info.items():
        if not val:
            continue
        row_data = [[Paragraph(field, label_style), Paragraph(val, value_style)]]
        t = Table(row_data, colWidths=[120, 350])
        t.setStyle(TableStyle([
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("TOPPADDING", (0,0), (-1,-1), 6),
            ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ("LINEBELOW", (0,0), (-1,-1), 0.3, colors.HexColor("#e5e7eb")),
        ]))
        story.append(t)

    if nganh_nghe:
        story.append(Spacer(1, 14))
        story.append(Paragraph("Ngành nghề kinh doanh", ParagraphStyle(
            "nn", fontName=font_name, fontSize=11, spaceAfter=6,
            textColor=colors.HexColor("#1a56db"))))
        data = [[Paragraph("<b>Mã</b>", cell_style), Paragraph("<b>Ngành</b>", cell_style)]]
        for nn in nganh_nghe:
            data.append([Paragraph(nn["Mã"], cell_style), Paragraph(nn["Ngành"], cell_style)])
        tbl = Table(data, colWidths=[60, 410], repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1a56db")),
            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
            ("GRID",       (0,0), (-1,-1), 0.3, colors.HexColor("#e5e7eb")),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f9fafb")]),
            ("VALIGN",     (0,0), (-1,-1), "TOP"),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ]))
        story.append(tbl)

    doc.build(story)
    return buf.getvalue()


with tab_lookup:
    st.subheader("🔎 Tra cứu thông tin doanh nghiệp theo MST")
    st.caption("Nhập mã số thuế để lấy toàn bộ thông tin từ vntax.net")

    import streamlit.components.v1 as components

    ICONS = {
        "Mã số thuế":     "🧾",
        "Địa chỉ thuế":   "📍",
        "Địa chỉ":        "📌",
        "Tình trạng":     "✅",
        "Người đại diện": "👤",
        "Điện thoại":     "📞",
        "Ngày hoạt động": "📅",
        "Quản lý bởi":    "🏛️",
        "Loại hình DN":   "🏢",
        "Ngành":          "🏭",
        "Tên quốc tế":    "🌐",
        "Tên viết tắt":   "✏️",
    }

    def build_html_card(ten, info, nganh_nghe):
        rows_html = ""
        for field, val in info.items():
            if not val:
                continue
            icon = ICONS.get(field, "•")
            if field == "Tình trạng":
                val_style = "color:#16a34a;font-weight:600;"
            elif field == "Người đại diện":
                val_style = "color:#1a56db;font-weight:600;"
            elif field == "Mã số thuế":
                val_style = "font-weight:600;"
            else:
                val_style = ""
            rows_html += f"""<div style="display:grid;grid-template-columns:180px 1fr;padding:10px 16px;border-bottom:1px solid #f3f4f6;gap:8px;align-items:start;">
  <div style="color:#6b7280;font-size:13px;display:flex;align-items:center;gap:6px;">{icon} {field}</div>
  <div style="font-size:14px;{val_style}">{val}</div>
</div>"""

        nganh_table = ""
        if nganh_nghe:
            rows_nn = ""
            for i, nn in enumerate(nganh_nghe):
                bg = "#f9fafb" if i % 2 else "#ffffff"
                rows_nn += f"""<tr style="background:{bg};">
  <td style="padding:6px 10px;border:1px solid #e5e7eb;font-size:13px;">{nn['Mã']}</td>
  <td style="padding:6px 10px;border:1px solid #e5e7eb;font-size:13px;">{nn['Ngành']}</td>
</tr>"""
            nganh_table = f"""<div style="padding:16px;">
  <p style="font-weight:600;font-size:14px;margin:0 0 8px;">Ngành nghề kinh doanh</p>
  <table style="width:100%;border-collapse:collapse;">
    <thead><tr style="background:#1a56db;color:white;">
      <th style="padding:8px 10px;border:1px solid #1a56db;font-size:13px;width:70px;text-align:left;">Mã</th>
      <th style="padding:8px 10px;border:1px solid #1a56db;font-size:13px;text-align:left;">Ngành</th>
    </tr></thead>
    <tbody>{rows_nn}</tbody>
  </table>
</div>"""

        return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"></head><body style="margin:0;padding:8px;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;">
<div style="background:white;border-radius:8px;box-shadow:0 1px 3px rgba(0,0,0,0.12);border:1px solid #e5e7eb;overflow:hidden;">
  <div style="padding:14px 16px;border-bottom:1px solid #e5e7eb;">
    <h1 style="font-size:15px;font-weight:600;color:#1a56db;text-transform:uppercase;margin:0;line-height:1.4;">{ten}</h1>
  </div>
  <div>{rows_html}</div>
  {nganh_table}
</div>
</body></html>"""

    col_ta, col_cfg = st.columns([3, 1])
    with col_ta:
        mst_input = st.text_area(
            "Mã số thuế (mỗi dòng 1 mã)",
            placeholder="VD:\n0901101740-003\n0100109106\n1234567890",
            height=120,
            label_visibility="collapsed",
        )
    with col_cfg:
        lookup_threads = st.slider("Số luồng", min_value=1, max_value=5, value=3)
        btn_lookup = st.button("Tra cứu", type="primary", use_container_width=True)

    if btn_lookup and mst_input.strip():
        mst_list = [m.strip() for m in mst_input.strip().splitlines() if m.strip()]
        total = len(mst_list)

        if total == 1:
            # Chế độ đơn: hiển thị card + tải PDF riêng
            with st.spinner("Đang tra cứu..."):
                result, err = lookup_mst(mst_list[0])
            if err:
                st.error(err)
            elif result:
                ten, info, nganh_nghe = result["ten"], result["info"], result["nganh_nghe"]
                html_card = build_html_card(ten, info, nganh_nghe)
                height = 120 + len([v for v in info.values() if v]) * 42 + len(nganh_nghe) * 34 + (80 if nganh_nghe else 0)
                components.html(html_card, height=height, scrolling=True)
                st.divider()
                mst_safe = mst_list[0].replace("/", "-")
                st.download_button(
                    "⬇ Tải PDF",
                    data=lookup_to_pdf(ten, info, nganh_nghe),
                    file_name=f"{mst_safe}.pdf",
                    mime="application/pdf",
                    key="dl_lookup_pdf",
                )
        else:
            # Chế độ nhiều MST: song song → gom ZIP
            log_placeholder = st.empty()
            progress_bar = st.progress(0)
            errors = []
            pdf_results = {}  # mst -> bytes
            done_count = 0

            with ThreadPoolExecutor(max_workers=lookup_threads) as ex:
                future_to_mst = {ex.submit(lookup_mst, mst): mst for mst in mst_list}
                for future in as_completed(future_to_mst):
                    mst = future_to_mst[future]
                    done_count += 1
                    log_placeholder.info(f"Đang xử lý: **{mst}** ({done_count}/{total})...")
                    progress_bar.progress(done_count / total)
                    result, err = future.result()
                    if err or not result:
                        errors.append(f"{mst}: {err or 'Không tìm thấy'}")
                        continue
                    ten, info, nganh_nghe = result["ten"], result["info"], result["nganh_nghe"]
                    pdf_results[mst] = lookup_to_pdf(ten, info, nganh_nghe)

            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for mst, pdf_bytes in pdf_results.items():
                    zf.writestr(f"{mst.replace('/', '-')}.pdf", pdf_bytes)

            log_placeholder.success(f"✅ Hoàn thành {len(pdf_results)}/{total} mã.")
            if errors:
                st.warning("Lỗi các mã:\n" + "\n".join(errors))

            st.download_button(
                f"⬇ Tải ZIP ({len(pdf_results)} file PDF)",
                data=zip_buf.getvalue(),
                file_name="mst_lookup.zip",
                mime="application/zip",
                key="dl_lookup_zip",
            )
